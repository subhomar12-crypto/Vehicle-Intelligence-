"""
PREDICT - Vehicle Intelligence Platform
Copyright © 2026 PREDICT
All rights reserved.

This file is proprietary and confidential.
Unauthorized copying, modification, distribution, or use is strictly prohibited.

Created: January 2026
Module: Create Executive Summary Docx

Generate Executive Summary Word Document for Investors
Company: Predict
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_executive_summary():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Executive Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Predict - AI-Powered Predictive Vehicle Maintenance System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 51, 102)

    # Document info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Prepared for: ').bold = True
    info.add_run('Investors and Strategic Partners\n')
    info.add_run('Date: ').bold = True
    info.add_run(f'{datetime.now().strftime("%B %Y")}\n')
    info.add_run('Version: ').bold = True
    info.add_run('2.0')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Problem Statement',
        '2. Solution Overview',
        '3. Market Opportunity',
        '4. Revenue Model',
        '5. Competitive Analysis',
        '6. Team',
        '7. Financial Projections',
        '8. Funding Ask',
        '9. System Overview',
        '10. Predict Guardian - Family & Fleet Monitoring App',
        '11. Benefits for Individual Vehicle Owners',
        '12. Benefits for Fleet Owners & Enterprises',
        '13. Hardware Requirements & Pricing',
        '14. PC Hardware Recommendations',
        '15. Startup Budget Estimate',
        '16. Competitive Advantages in Qatar & MENA',
        '17. Future R&D Roadmap'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # =========================================================================
    # Section 1: Problem Statement
    # =========================================================================
    doc.add_heading('1. Problem Statement', level=1)

    doc.add_heading('1.1 The $210 Billion Problem', level=2)
    doc.add_paragraph(
        'Vehicle breakdowns cost the global economy over $210 billion annually in emergency repairs, '
        'towing services, lost productivity, and secondary accidents. In the MENA region alone, '
        'extreme temperatures accelerate component wear, leading to 40% higher failure rates than '
        'temperate climates.'
    )

    doc.add_heading('1.2 Current Market Pain Points', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ['Pain Point', 'Impact', 'Who Suffers']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'B22222')  # Dark red for problem
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    pain_data = [
        ['Unexpected Breakdowns', 'QAR 500-5,000 per incident', 'Individual owners, families'],
        ['Fleet Downtime', '15-20% lost productivity', 'Taxi companies, delivery services'],
        ['Reactive Maintenance', '2-3x higher repair costs', 'All vehicle owners'],
        ['No Advance Warning', 'Roadside emergencies, safety risks', 'Drivers, passengers'],
        ['Climate-Accelerated Wear', '40% faster battery/AC failures', 'Qatar & MENA residents']
    ]
    for i, row_data in enumerate(pain_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('1.3 Why Existing Solutions Fail', level=2)
    failures = [
        'OBD-II Scanners: Only detect existing faults, no prediction capability',
        'Dealer Diagnostics: Expensive ($100-500), require visits, still reactive',
        'Fleet Telematics: Focus on GPS/tracking, limited maintenance intelligence',
        'Generic AI Solutions: Not trained on regional climate conditions',
        'Manual Inspections: Time-consuming, subjective, miss hidden degradation'
    ]
    for failure in failures:
        doc.add_paragraph(failure, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('The Gap: ').bold = True
    p.add_run('No affordable, consumer-accessible solution predicts failures BEFORE they happen.')

    doc.add_page_break()

    # =========================================================================
    # Section 2: Solution Overview
    # =========================================================================
    doc.add_heading('2. Solution Overview', level=1)

    doc.add_heading('2.1 Predict: AI That Sees Failures Coming', level=2)
    doc.add_paragraph(
        'Predict is a comprehensive AI-powered predictive maintenance platform designed to anticipate vehicle '
        'failures 7-60 days before they occur. Using deep learning (LSTM neural networks) that will be trained on '
        'real vehicle telemetry, the system aims to transform reactive maintenance into proactive prevention.'
    )

    # Value proposition box
    p = doc.add_paragraph()
    p.add_run('One Sentence: ').bold = True
    p.add_run('"Predict tells you what will break in your car next month, so you can fix it this weekend."')

    doc.add_heading('2.2 How It Works', level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['Step', 'What Happens']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '006400')  # Dark green for solution
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    how_data = [
        ['1. Connect', 'Plug OBD-II adapter into vehicle, pair with smartphone app'],
        ['2. Collect', 'System continuously gathers 50+ vehicle parameters while driving'],
        ['3. Analyze', 'LSTM neural network detects degradation patterns in real-time'],
        ['4. Predict', 'AI identifies which component will fail and when (target: 7-60 days ahead)']
    ]
    for i, row_data in enumerate(how_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('2.3 Platform Components', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Component', 'Platform', 'Function']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    platform_data = [
        ['Predict Desktop', 'Windows PC', 'AI training hub, fleet management, analytics'],
        ['Predict App', 'Android', 'Real-time monitoring, driver alerts, OBD-II collection'],
        ['Predict Guardian', 'Android', 'Family oversight, driving behavior, emergency tracking'],
        ['Predict Cloud', 'FastAPI/Linux', 'API gateway, data sync, multi-device coordination']
    ]
    for i, row_data in enumerate(platform_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('2.4 Key Differentiators', level=2)
    differentiators = [
        '7-60 Day Prediction Window (Target): Aiming for longest advance warning in the market',
        'Deep Learning AI: LSTM neural networks, not simple threshold alerts',
        'Climate-Optimized: Designed for MENA extreme heat conditions',
        'Affordable Entry: Under QAR 1,500 per vehicle',
        'Continuous Learning: Improves with every mile driven'
    ]
    for diff in differentiators:
        doc.add_paragraph(diff, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 3: Market Opportunity
    # =========================================================================
    doc.add_heading('3. Market Opportunity', level=1)

    doc.add_heading('3.1 Total Addressable Market (TAM)', level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Market', 'Size', 'Value']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4169E1')  # Royal blue
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tam_data = [
        ['Global Vehicle Maintenance', '1.4 billion vehicles', '$750 billion/year'],
        ['Predictive Maintenance Software', '50 million potential users', '$15 billion/year'],
        ['Connected Car Services', '500 million connected vehicles by 2030', '$200 billion/year']
    ]
    for i, row_data in enumerate(tam_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('3.2 Serviceable Addressable Market (SAM)', level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4169E1')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    sam_data = [
        ['MENA Region Vehicles', '50 million vehicles', '$15 billion maintenance market'],
        ['MENA Fleet Vehicles', '5 million commercial vehicles', '$3 billion fleet maintenance'],
        ['GCC Countries', '12 million vehicles', '$4 billion market']
    ]
    for i, row_data in enumerate(sam_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('3.3 Serviceable Obtainable Market (SOM)', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Region', 'Target (5 Years)', 'Revenue Potential']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4169E1')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    som_data = [
        ['Qatar', '25,000 vehicles (2% penetration)', 'QAR 50 million ARR'],
        ['UAE', '50,000 vehicles', 'QAR 100 million ARR'],
        ['Saudi Arabia', '100,000 vehicles', 'QAR 200 million ARR'],
        ['Total GCC (Year 5)', '175,000 vehicles', 'QAR 350 million ARR']
    ]
    for i, row_data in enumerate(som_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('3.4 Market Drivers', level=2)
    drivers = [
        'Connected Vehicle Growth: 30% annual increase in OBD-II capable vehicles',
        'Insurance Integration: Insurers offering discounts for predictive maintenance users',
        'Extreme Climate Impact: MENA heat accelerates vehicle wear, increasing demand',
        'EV Transition: Battery health prediction becomes critical for EV adoption',
        'Regulatory Push: Qatar/UAE governments mandating fleet efficiency standards'
    ]
    for driver in drivers:
        doc.add_paragraph(driver, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 4: Revenue Model
    # =========================================================================
    doc.add_heading('4. Revenue Model', level=1)

    # Important note about bundled model
    p = doc.add_paragraph()
    p.add_run('BUNDLED HARDWARE MODEL: ').bold = True
    p.add_run(
        'All customers must purchase our proprietary hardware kit (OBD-II adapter + dual ESP32 sensors). '
        'This ensures data quality for AI predictions and creates a hardware revenue stream. '
        'No BYOD (bring-your-own-device) option - our sensors are required for accurate predictions.'
    )

    doc.add_paragraph()

    doc.add_heading('4.1 Pricing Structure', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Item', 'Price (QAR)', 'Type', 'Margin']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '228B22')  # Forest green
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    revenue_data = [
        ['Hardware Kit (required)', '450', 'One-time purchase', '40%'],
        ['Annual Subscription', '500/year', 'Recurring (annual only)', '90%'],
        ['Fleet Hardware (10+ vehicles)', '350/vehicle', 'One-time (volume discount)', '35%'],
        ['Fleet Subscription (10+ vehicles)', '400/vehicle/year', 'Recurring (volume discount)', '90%']
    ]
    for i, row_data in enumerate(revenue_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('4.2 Customer Pricing Summary', level=2)

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'

    headers = ['Customer Type', 'Hardware', 'Year 1 Total', 'Year 2+ (Renewal)', 'Per Day Cost']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '228B22')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tier_data = [
        ['Individual (1 car)', 'QAR 450', 'QAR 950', 'QAR 500/year', 'QAR 1.37/day'],
        ['Family (3 cars)', 'QAR 1,350', 'QAR 2,850', 'QAR 1,500/year', 'QAR 1.37/car/day'],
        ['Fleet (10 vehicles)', 'QAR 3,500', 'QAR 7,500', 'QAR 4,000/year', 'QAR 1.10/car/day']
    ]
    for i, row_data in enumerate(tier_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Why Annual Only (No Monthly): ').bold = True
    p.add_run(
        'Annual subscriptions ensure customer commitment, reduce churn, provide predictable revenue, '
        'and give us time to collect enough driving data for accurate AI predictions. '
        'Monthly subscriptions would encourage customers to cancel before seeing the full value.'
    )

    doc.add_paragraph()

    doc.add_heading('4.3 Unit Economics', level=2)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    headers = ['Metric', 'Value']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '228B22')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    economics_data = [
        ['Hardware Cost (to us)', 'QAR 270 (~$75)'],
        ['Hardware Sale Price', 'QAR 450'],
        ['Hardware Margin', 'QAR 180 (40%)'],
        ['Annual Subscription Revenue', 'QAR 500'],
        ['Customer Year 1 Value', 'QAR 680 (hardware margin + subscription)'],
        ['3-Year LTV (if renewed)', 'QAR 1,180 per customer']
    ]
    for i, row_data in enumerate(economics_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_page_break()

    # =========================================================================
    # Section 5: Competitive Analysis
    # =========================================================================
    doc.add_heading('5. Competitive Analysis', level=1)

    doc.add_heading('5.1 Competitive Landscape', level=2)

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    headers = ['Competitor', 'Type', 'Prediction?', 'Price', 'Weakness']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'DC143C')  # Crimson
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ['OBD-II Scanners', 'Hardware', 'No', '$20-100', 'Only reads existing faults'],
        ['Dealer Service', 'Service', 'No', '$100-500/visit', 'Reactive, expensive, time-consuming'],
        ['Fleet Telematics', 'Software', 'Limited', '$30-50/month', 'GPS focus, basic maintenance'],
        ['Zubie/Bouncie', 'Consumer', 'Basic', '$10-15/month', 'Simple alerts, no AI prediction'],
        ['Predict', 'AI Platform', '7-60 days*', 'QAR 950 Year 1*', '*Hardware + annual subscription bundle']
    ]
    for i, row_data in enumerate(comp_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 4:  # Predict row
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i + 1].cells[j], 'E6FFE6')

    doc.add_paragraph()

    doc.add_heading('5.2 Why Predict Wins', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Advantage', 'Predict', 'Competitors']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '006400')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    win_data = [
        ['Prediction Window', '7-60 days advance (target)', 'Real-time or none'],
        ['AI Technology', 'LSTM deep learning', 'Rule-based thresholds'],
        ['Climate Optimization', 'MENA heat-trained', 'Generic global models'],
        ['Price Point', 'QAR 500/year + hardware', '$200-500/month (enterprise only)']
    ]
    for i, row_data in enumerate(win_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('5.3 Barriers to Entry', level=2)
    barriers = [
        'Proprietary AI Models: 2+ years of training data and algorithm refinement',
        'Regional Expertise: Climate-specific failure patterns not available to global players',
        'Hardware Integration: Dual ESP32 sensor system with custom firmware',
        'First Mover Advantage: Building network effects in Qatar/GCC market',
        'Regulatory Relationships: Working with local transport authorities'
    ]
    for barrier in barriers:
        doc.add_paragraph(barrier, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 6: Team
    # =========================================================================
    doc.add_heading('6. Team', level=1)

    doc.add_heading('6.1 Core Team', level=2)
    doc.add_paragraph(
        'The Predict team combines deep expertise in AI/ML, automotive systems, and regional market knowledge.'
    )

    # Team table with placeholders
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Role', 'Expertise', 'Background']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4B0082')  # Indigo
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    team_data = [
        ['Founder / CEO', 'Product Vision, Business Strategy', 'Tech entrepreneur, Qatar-based'],
        ['Technical Lead', 'AI/ML, LSTM Neural Networks, Python', 'Deep learning specialist'],
        ['Hardware Lead', 'ESP32, IoT, OBD-II Integration', 'Embedded systems engineer']
    ]
    for i, row_data in enumerate(team_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('6.2 Advisory Board (Planned)', level=2)
    advisors = [
        'Automotive Industry Expert: Former OEM executive with MENA market experience',
        'AI/ML Advisor: Academic or industry researcher in predictive analytics',
        'Fleet Management Expert: Senior executive from regional logistics company',
        'Insurance Industry Advisor: Underwriting or product leader from GCC insurer'
    ]
    for advisor in advisors:
        doc.add_paragraph(advisor, style='List Bullet')

    doc.add_heading('6.3 Hiring Plan', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Phase', 'Hires', 'Focus']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4B0082')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    hire_data = [
        ['Seed (0-12 mo)', '3-5 engineers', 'Android dev, backend, data science'],
        ['Series A (12-24 mo)', '10-15 team', 'Sales, customer success, ML ops'],
        ['Series B (24-36 mo)', '25-40 team', 'Regional expansion, enterprise sales'],
        ['Scale (36+ mo)', '50+ team', 'International markets, R&D lab']
    ]
    for i, row_data in enumerate(hire_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_page_break()

    # =========================================================================
    # Section 7: Financial Projections (Conservative/Honest)
    # =========================================================================
    doc.add_heading('7. Financial Projections', level=1)

    # Important disclaimer
    p = doc.add_paragraph()
    p.add_run('IMPORTANT DISCLAIMER: ').bold = True
    p.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    p.add_run(
        'These projections are ESTIMATES based on assumptions that are NOT yet validated. '
        'We have zero paying customers and zero revenue today. Numbers assume successful product '
        'development, market acceptance, and execution - none of which is guaranteed.'
    )

    doc.add_paragraph()

    doc.add_heading('7.1 Revenue Scenarios (QAR Thousands)', level=2)

    doc.add_paragraph(
        'We present three scenarios: Conservative (things go slower than expected), '
        'Base Case (reasonable execution), and Optimistic (everything works well). '
        'Investors should plan for Conservative scenario.'
    )

    doc.add_paragraph()

    # Conservative Scenario
    p = doc.add_paragraph()
    p.add_run('Conservative Scenario (40% probability):').bold = True

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Metric', 'Year 1', 'Year 2', 'Year 3']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'B22222')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    conservative_data = [
        ['Connected Vehicles', '30', '150', '400'],
        ['Paying Customers', '2', '8', '25'],
        ['Monthly Revenue (end of year)', 'QAR 3K', 'QAR 15K', 'QAR 50K'],
        ['Annual Revenue', 'QAR 18K', 'QAR 108K', 'QAR 450K']
    ]
    for i, row_data in enumerate(conservative_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    # Base Case
    p = doc.add_paragraph()
    p.add_run('Base Case Scenario (40% probability):').bold = True

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'DAA520')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    base_data = [
        ['Connected Vehicles', '100', '400', '1,000'],
        ['Paying Customers', '5', '20', '60'],
        ['Monthly Revenue (end of year)', 'QAR 10K', 'QAR 50K', 'QAR 150K'],
        ['Annual Revenue', 'QAR 60K', 'QAR 360K', 'QAR 1.2M']
    ]
    for i, row_data in enumerate(base_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    # Optimistic
    p = doc.add_paragraph()
    p.add_run('Optimistic Scenario (20% probability):').bold = True

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '228B22')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    optimistic_data = [
        ['Connected Vehicles', '200', '800', '2,500'],
        ['Paying Customers', '10', '50', '150'],
        ['Monthly Revenue (end of year)', 'QAR 25K', 'QAR 125K', 'QAR 400K'],
        ['Annual Revenue', 'QAR 150K', 'QAR 900K', 'QAR 3.6M']
    ]
    for i, row_data in enumerate(optimistic_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('7.2 Revenue Model Assumptions (Bundled Annual Model)', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Revenue Type', 'Pricing (QAR)', 'Notes']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    pricing_data = [
        ['Hardware Kit (required)', '450 one-time', 'OBD + dual ESP32 sensors, 40% margin'],
        ['Annual Subscription', '500/year', 'Annual only, no monthly option'],
        ['Fleet Hardware (10+)', '350/vehicle', 'Volume discount'],
        ['Fleet Subscription (10+)', '400/vehicle/year', 'Volume discount']
    ]
    for i, row_data in enumerate(pricing_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Note: ').italic = True
    p.add_run('These prices are ASSUMPTIONS. We have not tested willingness to pay. '
              'Hardware requirement may limit initial adoption but ensures data quality.')

    doc.add_paragraph()

    doc.add_heading('7.3 Cost Structure (Monthly Burn Rate)', level=2)

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    headers = ['Cost Category', 'Minimum Team', 'Growth Team', 'Full Team']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    cost_data = [
        ['Salaries (team)', 'QAR 16K', 'QAR 40K', 'QAR 80K'],
        ['Cloud/Infrastructure', 'QAR 1K', 'QAR 3K', 'QAR 8K'],
        ['Office/Misc', 'QAR 2K', 'QAR 5K', 'QAR 10K'],
        ['Marketing', 'QAR 1K', 'QAR 5K', 'QAR 15K'],
        ['Hardware Inventory', 'QAR 2K', 'QAR 5K', 'QAR 10K'],
        ['TOTAL MONTHLY BURN', 'QAR 22K', 'QAR 58K', 'QAR 123K']
    ]
    for i, row_data in enumerate(cost_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 5:  # Total row
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('7.4 Path to Break-Even', level=2)

    doc.add_paragraph(
        'Break-even requires enough paying customers to cover monthly costs. '
        'Assuming QAR 58K/month burn rate (growth team) and annual subscription model:'
    )

    breakeven_calc = [
        'Annual burn rate: QAR 58K x 12 = QAR 696K/year',
        'Revenue per new customer (Year 1): QAR 950 (hardware + subscription)',
        'Revenue per renewal: QAR 500/year',
        'Break-even (Year 1 only): ~735 new customers needed',
        'With 50% renewal rate: ~500 new + 250 renewals = sustainable',
        'Realistic timeline to break-even: 24-36 months after launch'
    ]
    for calc in breakeven_calc:
        doc.add_paragraph(calc, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Honest Assessment: ').bold = True
    p.add_run(
        'We will likely need multiple funding rounds before break-even. '
        'This seed investment is to prove the concept, not to reach profitability.'
    )

    doc.add_page_break()

    # =========================================================================
    # Section 8: Funding Ask
    # =========================================================================
    doc.add_heading('8. Funding Ask', level=1)

    # Current State - Honest Assessment
    doc.add_heading('8.1 Current Product State (Honest Assessment)', level=2)

    p = doc.add_paragraph()
    p.add_run('What We Have Built:').bold = True

    current_state = [
        'Desktop Application: Fully functional Windows app with AI prediction engine (complete)',
        'LSTM Neural Network: Deep learning model architecture implemented, trained on synthetic data only',
        'Cloud Server: FastAPI backend with 17 API endpoints ready for Android connection',
        'Vehicle Profile System: Multi-vehicle management with subscription/API key system',
        'Enterprise Features: Backup system, encryption, monitoring alerts implemented'
    ]
    for item in current_state:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('What We Need (The Gap):').bold = True

    gaps = [
        'Android App: Not yet developed - required for OBD-II data collection from vehicles',
        'Real Vehicle Data: Zero real-world OBD-II data collected - AI currently uses synthetic data',
        'Model Validation: Prediction accuracy unproven with actual vehicle failures',
        'Paying Customers: Zero revenue, zero subscriptions, zero fleet customers',
        'Hardware Testing: ESP32 sensor integration designed but not field-tested'
    ]
    for item in gaps:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # The Problem We Solve
    doc.add_heading('8.2 The Problem We Solve', level=2)

    p = doc.add_paragraph()
    p.add_run('Real-World Scenario:').bold = True
    doc.add_paragraph(
        'A taxi driver in Doha is driving passengers when his car battery suddenly dies in 45°C heat. '
        'He loses 4 hours of work (QAR 200-400 lost income), pays QAR 150 for towing, and QAR 400 for '
        'emergency battery replacement. Total cost: QAR 750-950 for ONE incident that could have been '
        'predicted 2-3 weeks earlier with proper monitoring.'
    )

    doc.add_paragraph()

    # Cost breakdown table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    headers = ['Cost Type', 'Without Predict', 'With Predict']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'B22222')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    problem_data = [
        ['Lost work time', 'QAR 200-400', 'QAR 0'],
        ['Towing service', 'QAR 100-200', 'QAR 0'],
        ['Emergency repair premium', '50-100% markup', 'Normal price'],
        ['Stress & inconvenience', 'High', 'Planned service'],
        ['Safety risk', 'Stranded in heat', 'None'],
        ['TOTAL per incident', 'QAR 500-1,500+', 'QAR 500/year subscription']
    ]
    for i, row_data in enumerate(problem_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 5:
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True
                if j == 1:
                    set_cell_shading(table.rows[i + 1].cells[j], 'FFE6E6')
                elif j == 2:
                    set_cell_shading(table.rows[i + 1].cells[j], 'E6FFE6')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('The Math: ').bold = True
    p.add_run(
        'If a fleet of 50 taxis experiences just 2 breakdowns per vehicle per year (conservative), '
        'that is 100 incidents x QAR 750 = QAR 75,000/year in preventable losses. Our solution at '
        'QAR 750/vehicle Year 1 (hardware + subscription) = QAR 37,500. Net savings: QAR 37,500/year (50%). '
        'Year 2+: Only QAR 400/vehicle/year renewal = QAR 20,000. Savings: QAR 55,000/year (73%).'
    )

    doc.add_paragraph()

    # Investment Options with Honest Numbers
    doc.add_heading('8.3 Investment Options', level=2)

    doc.add_paragraph(
        'We offer four investment tiers based on your risk appetite. All projections assume successful '
        'Android app development and real-world validation - which is NOT yet proven.'
    )

    doc.add_paragraph()

    # Option A: Micro Seed
    p = doc.add_paragraph()
    p.add_run('Option A: Micro Seed (Ultra Budget)').bold = True
    p.runs[0].font.color.rgb = RGBColor(70, 130, 180)

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ['Investment', 'Equity', 'Pre-Money Valuation', 'Runway']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4682B4')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    micro_data = ['QAR 100,000', '8-10%', 'QAR 1-1.25 Million', '4-6 months']
    for j, cell_data in enumerate(micro_data):
        table.rows[1].cells[j].text = cell_data
        table.rows[1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Use of Funds (QAR 100,000):').bold = True

    micro_funds = doc.add_table(rows=6, cols=3)
    micro_funds.style = 'Table Grid'
    micro_headers = ['Item', 'Amount (QAR)', 'Purpose']
    for i, h in enumerate(micro_headers):
        micro_funds.rows[0].cells[i].text = h
        micro_funds.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(micro_funds.rows[0].cells[i], '4682B4')
        micro_funds.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    micro_fund_data = [
        ['Android Developer (contract)', '40,000', '3-4 months to build MVP app'],
        ['OBD-II Hardware (20 units)', '4,000', 'Testing devices for pilot vehicles'],
        ['Cloud Hosting (6 months)', '6,000', 'Server costs for API and data storage'],
        ['Founder Salary (minimal)', '36,000', 'QAR 6,000/month x 6 months'],
        ['Miscellaneous', '14,000', 'Legal, testing, contingency']
    ]
    for i, row_data in enumerate(micro_fund_data):
        for j, cell_data in enumerate(row_data):
            micro_funds.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Milestones for QAR 100K:').bold = True
    micro_milestones = [
        'Android app MVP (basic OBD-II connection + data upload)',
        '10-20 test vehicles collecting real data',
        'First validation: Does predicted failure match reality?',
        'Decision point: Continue or pivot based on results'
    ]
    for m in micro_milestones:
        doc.add_paragraph(m, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Risk Level: HIGHEST').bold = True
    p.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    p.add_run(' - Unproven product, no revenue, no validation yet')

    doc.add_paragraph()

    # Option B: Budget
    p = doc.add_paragraph()
    p.add_run('Option B: Seed Round (Budget Friendly)').bold = True
    p.runs[0].font.color.rgb = RGBColor(34, 139, 34)

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ['Investment', 'Equity', 'Pre-Money Valuation', 'Runway']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '228B22')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    budget_data = ['QAR 300,000 - 400,000', '12-15%', 'QAR 2-2.5 Million', '10-14 months']
    for j, cell_data in enumerate(budget_data):
        table.rows[1].cells[j].text = cell_data
        table.rows[1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Use of Funds (QAR 350,000):').bold = True

    budget_funds = doc.add_table(rows=7, cols=3)
    budget_funds.style = 'Table Grid'
    for i, h in enumerate(micro_headers):
        budget_funds.rows[0].cells[i].text = h
        budget_funds.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(budget_funds.rows[0].cells[i], '228B22')
        budget_funds.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    budget_fund_data = [
        ['Android Developer (full-time)', '120,000', 'QAR 10K/month x 12 months'],
        ['Backend/AI Developer (part-time)', '60,000', 'QAR 5K/month x 12 months'],
        ['OBD-II Hardware (50 units)', '10,000', 'Pilot fleet hardware'],
        ['Cloud Infrastructure', '24,000', 'QAR 2K/month x 12 months'],
        ['Founder Salary', '96,000', 'QAR 8K/month x 12 months'],
        ['Marketing + Operations', '40,000', 'Pilot customer acquisition, legal']
    ]
    for i, row_data in enumerate(budget_fund_data):
        for j, cell_data in enumerate(row_data):
            budget_funds.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Milestones for QAR 350K:').bold = True
    budget_milestones = [
        'Full Android app on Google Play Store',
        '50-100 vehicles with 6+ months of data',
        '2-3 paying pilot fleet customers (proof of willingness to pay)',
        'AI model retrained on real data with measured accuracy',
        'First revenue: Target QAR 5,000-15,000/month by month 12'
    ]
    for m in budget_milestones:
        doc.add_paragraph(m, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Risk Level: HIGH').bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 140, 0)
    p.add_run(' - Product being validated, early revenue possible')

    doc.add_paragraph()

    # Option C: Standard
    p = doc.add_paragraph()
    p.add_run('Option C: Seed Round (Standard)').bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 140, 0)

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ['Investment', 'Equity', 'Pre-Money Valuation', 'Runway']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'FF8C00')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    standard_data = ['QAR 600,000 - 800,000', '15-20%', 'QAR 3-4 Million', '18-24 months']
    for j, cell_data in enumerate(standard_data):
        table.rows[1].cells[j].text = cell_data
        table.rows[1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Use of Funds (QAR 700,000):').bold = True

    standard_funds = doc.add_table(rows=8, cols=3)
    standard_funds.style = 'Table Grid'
    for i, h in enumerate(micro_headers):
        standard_funds.rows[0].cells[i].text = h
        standard_funds.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(standard_funds.rows[0].cells[i], 'FF8C00')
        standard_funds.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    standard_fund_data = [
        ['Android Developer (full-time)', '216,000', 'QAR 12K/month x 18 months'],
        ['Backend Developer (full-time)', '180,000', 'QAR 10K/month x 18 months'],
        ['Sales/BD Person', '108,000', 'QAR 6K/month x 18 months (starts month 6)'],
        ['Hardware (200 units)', '40,000', 'Fleet deployment hardware'],
        ['Cloud Infrastructure', '54,000', 'QAR 3K/month x 18 months (scaling)'],
        ['Founder Salary', '162,000', 'QAR 9K/month x 18 months'],
        ['Marketing + Legal + Office', '40,000', 'Customer acquisition, compliance']
    ]
    for i, row_data in enumerate(standard_fund_data):
        for j, cell_data in enumerate(row_data):
            standard_funds.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Milestones for QAR 700K:').bold = True
    standard_milestones = [
        'Complete platform: Android + Guardian + Desktop + Cloud',
        '200-500 connected vehicles',
        '5-10 paying fleet customers',
        'Monthly recurring revenue: QAR 20,000-50,000 by month 18',
        'Validated prediction accuracy: Target 70%+ correct predictions',
        'Ready for Series A fundraising'
    ]
    for m in standard_milestones:
        doc.add_paragraph(m, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Risk Level: MODERATE').bold = True
    p.runs[0].font.color.rgb = RGBColor(218, 165, 32)
    p.add_run(' - Longer runway allows for iteration and market learning')

    doc.add_paragraph()

    # Option D: Accelerated
    p = doc.add_paragraph()
    p.add_run('Option D: Seed Round (Accelerated Growth)').bold = True
    p.runs[0].font.color.rgb = RGBColor(139, 0, 0)

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ['Investment', 'Equity', 'Pre-Money Valuation', 'Runway']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '8B0000')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    large_data = ['QAR 1.5 Million+', '20-25%', 'QAR 5-6 Million', '24-30 months']
    for j, cell_data in enumerate(large_data):
        table.rows[1].cells[j].text = cell_data
        table.rows[1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Use of Funds (QAR 1,500,000):').bold = True

    large_funds = doc.add_table(rows=9, cols=3)
    large_funds.style = 'Table Grid'
    for i, h in enumerate(micro_headers):
        large_funds.rows[0].cells[i].text = h
        large_funds.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(large_funds.rows[0].cells[i], '8B0000')
        large_funds.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    large_fund_data = [
        ['Engineering Team (4 people)', '720,000', 'Android, Backend, AI, QA x 24 months'],
        ['Sales & Customer Success (2)', '288,000', 'BD + Support x 24 months'],
        ['Hardware (500+ units)', '100,000', 'Fleet deployment at scale'],
        ['Cloud Infrastructure', '120,000', 'QAR 5K/month x 24 months'],
        ['Office Space', '72,000', 'QAR 3K/month x 24 months'],
        ['Marketing & Partnerships', '80,000', 'Events, insurance partnerships'],
        ['Founder Salary', '240,000', 'QAR 10K/month x 24 months'],
        ['Legal, Compliance, Buffer', '80,000', 'Contingency and operations']
    ]
    for i, row_data in enumerate(large_fund_data):
        for j, cell_data in enumerate(row_data):
            large_funds.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Milestones for QAR 1.5M:').bold = True
    large_milestones = [
        '1,000+ connected vehicles across Qatar',
        '15-25 paying fleet/enterprise customers',
        'Monthly recurring revenue: QAR 75,000-150,000 by month 24',
        'UAE market entry (pilot customers)',
        '1 insurance company partnership (usage-based insurance)',
        'Series A ready with proven metrics'
    ]
    for m in large_milestones:
        doc.add_paragraph(m, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Risk Level: LOWER').bold = True
    p.runs[0].font.color.rgb = RGBColor(34, 139, 34)
    p.add_run(' - Full team, long runway, room for pivots and learning')

    doc.add_paragraph()

    # Comparison Table
    doc.add_heading('8.4 Investment Comparison Summary', level=2)

    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'

    headers = ['Metric', 'Micro (100K)', 'Budget (350K)', 'Standard (700K)', 'Accelerated (1.5M)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'DAA520')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    comparison_data = [
        ['Runway', '4-6 months', '10-14 months', '18-24 months', '24-30 months'],
        ['Team Size', '1 + contractor', '2-3 people', '4-5 people', '7-8 people'],
        ['Target Vehicles', '10-20', '50-100', '200-500', '1,000+'],
        ['Target Customers', '0 (testing)', '2-3 pilots', '5-10 paying', '15-25 paying'],
        ['Target MRR (end)', 'QAR 0', 'QAR 5-15K', 'QAR 20-50K', 'QAR 75-150K'],
        ['Android App', 'Basic MVP', 'Full app', 'Full + Guardian', 'Full suite'],
        ['AI Validation', 'Initial test', 'Real data training', 'Proven accuracy', 'Production ready'],
        ['Risk Level', 'Highest', 'High', 'Moderate', 'Lower']
    ]
    for i, row_data in enumerate(comparison_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    # Honest Risk Disclosure
    doc.add_heading('8.5 Risk Disclosure', level=2)

    p = doc.add_paragraph()
    p.add_run('Key Risks Investors Should Understand:').bold = True

    risks = [
        'Technical Risk: LSTM prediction accuracy is unproven with real vehicle data. Model may need significant retraining.',
        'Market Risk: Qatar vehicle owners may not adopt mobile-based solutions. Fleet operators may resist change.',
        'Execution Risk: Android app development timeline may slip. Finding quality developers in Qatar is challenging.',
        'Competition Risk: Global players (Zubie, Bouncie) may enter GCC market with larger budgets.',
        'Revenue Risk: Subscription pricing assumptions (QAR 50-100/month) are untested. Customers may not pay expected rates.',
        'Data Risk: Collecting sufficient training data (thousands of vehicle-months) takes time. AI needs data to improve.'
    ]
    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('8.6 Why Invest Despite Risks?', level=2)
    why_invest = [
        'First Mover: No AI predictive maintenance solution exists in Qatar/GCC - we have head start',
        'Working Foundation: Desktop app, server, AI architecture all built and functional',
        'Climate Need: MENA extreme heat causes 40% higher vehicle failure rates - real problem to solve',
        'Low Competition: Regional players focus on GPS tracking, not prediction',
        'Founder Commitment: Technical founder with deep domain knowledge, not outsourcing core product',
        'Lean Approach: We are asking for realistic amounts with specific milestones, not inflated "vision" funding'
    ]
    for item in why_invest:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 9: System Overview
    # =========================================================================
    doc.add_heading('9. System Overview', level=1)

    doc.add_heading('9.1 Technical Architecture', level=2)
    doc.add_paragraph(
        'Predict is a comprehensive AI-powered predictive maintenance platform designed to anticipate '
        'vehicle failures 7-60 days before they occur. The system combines real-time OBD-II diagnostics, '
        'deep learning (LSTM neural networks), external IoT sensors, and a continuous feedback loop '
        'with the goal of delivering high accuracy in failure prediction.'
    )

    doc.add_paragraph('The platform consists of three integrated components:')

    # Components table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Component', 'Platform', 'Primary Function']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '003366')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    data = [
        ['Desktop Application', 'Windows PC', 'Central AI hub, LSTM training, fleet management, data analytics'],
        ['Predict App (Android)', 'Mobile/Tablet', 'Real-time monitoring, driver alerts, OBD-II data collection'],
        ['Predict Guardian', 'Mobile/Tablet', 'Family oversight, driving behavior, emergency location tracking'],
        ['Cloud Server', 'FastAPI/Linux', 'API gateway, data synchronization, multi-device coordination']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data

    doc.add_paragraph()

    # Desktop Application
    doc.add_heading('9.2 Desktop Application (Windows)', level=2)
    doc.add_paragraph('The Desktop Application serves as the command center for the entire system:')

    features = [
        'LSTM Deep Learning Engine: TensorFlow-powered neural network trained on vehicle telemetry patterns',
        'Multi-Profile Management: Support for unlimited vehicle profiles with individual AI models',
        'Advanced Analytics Dashboard: Real-time health scoring across 8 vehicle subsystems',
        'Service History Tracking: Complete maintenance records with AI learning integration',
        'Synthetic Data Generation: Bootstrap training capability for new deployments',
        'Enterprise Features: Automated backups, data encryption, monitoring alerts',
        'API Key Management: Secure key generation for Android app authentication'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Technology Stack: ').bold = True
    p.add_run('Python 3.11+, PySide6/PyQt5 GUI, TensorFlow 2.x / Keras, SQLite databases, FastAPI for local server')

    # Android Application
    doc.add_heading('9.3 Android Application', level=2)
    doc.add_paragraph('The Android Application provides mobile connectivity and real-time data collection:')

    android_features = [
        'Bluetooth OBD-II Connection: Direct link to vehicle diagnostic port via ELM327 adapter',
        'Real-Time Telemetry: Live display of 50+ OBD-II parameters',
        'AI Prediction Display: Shows failure predictions with confidence scores',
        'Push Notifications: Alerts for critical predictions and maintenance reminders',
        'Offline Capability: Local data storage with background sync',
        'ESP32 Sensor Integration: Receives data from external IoT sensors'
    ]
    for feature in android_features:
        doc.add_paragraph(feature, style='List Bullet')

    # Cloud Server
    doc.add_heading('9.4 Cloud Server (FastAPI)', level=2)
    doc.add_paragraph('The Cloud Server enables multi-device synchronization and remote access:')

    server_features = [
        'RESTful API: 17 endpoints covering all system functions',
        'Secure Authentication: API key-based with SHA-256 hashing',
        'Real-Time Processing: Connects to Desktop AI engine for predictions',
        'Database Management: SQLite with automated backups',
        'Cloudflare Tunnel: Secure HTTPS exposure without port forwarding'
    ]
    for feature in server_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 10: Predict Guardian - Family & Fleet Monitoring App
    # =========================================================================
    doc.add_heading('10. Predict Guardian - Family & Fleet Monitoring App', level=1)

    doc.add_paragraph(
        'Predict Guardian is a companion mobile application designed for family heads, fleet managers, '
        'and business owners who need oversight of multiple vehicles and drivers registered under a single account. '
        'This powerful monitoring solution provides peace of mind for parents, enhanced safety for families, '
        'and operational visibility for small business owners.'
    )

    doc.add_heading('10.1 Target Users', level=2)
    target_users = [
        'Parents monitoring teenage drivers',
        'Family heads overseeing household vehicles (spouse, children, elderly parents)',
        'Small business owners with company vehicles',
        'Private employers monitoring driver staff (chauffeurs, delivery personnel)',
        'Fleet supervisors requiring real-time oversight'
    ]
    for user in target_users:
        doc.add_paragraph(user, style='List Bullet')

    doc.add_heading('10.2 Core Features', level=2)

    # Features table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    headers = ['Feature', 'Description']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '6B8E23')  # Olive green for Guardian
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    guardian_features = [
        ['Multi-Vehicle Dashboard', 'View all family/fleet vehicles in one unified interface with real-time status'],
        ['Driving Behavior Analysis', 'Monitor speed patterns, harsh braking, rapid acceleration, and driving scores'],
        ['Emergency Location Request', 'Request real-time GPS location from any registered vehicle for emergencies'],
        ['Service Timeline Tracking', 'View upcoming and overdue maintenance for all vehicles in one place'],
        ['Driver Profiles', 'Assign vehicles to specific drivers (son, daughter, wife, driver) with individual tracking'],
        ['Alert Notifications', 'Receive instant alerts for speeding, geofence violations, or vehicle issues']
    ]
    for i, row_data in enumerate(guardian_features):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('10.3 Safety & Monitoring Capabilities', level=2)
    safety_caps = [
        'Real-Time Location Tracking: View current position of any vehicle on demand',
        'Geofencing: Set safe zones and receive alerts when vehicles enter/exit areas',
        'Speed Monitoring: Track speed violations with timestamp and location data',
        'Trip History: Review complete trip logs with routes, stops, and durations',
        'Driving Score: AI-generated safety scores based on driving patterns',
        'Fatigue Detection: Alerts for unusual driving patterns indicating driver fatigue',
        'Curfew Alerts: Notifications when vehicles are used outside designated hours'
    ]
    for cap in safety_caps:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_heading('10.4 Privacy & Consent Model', level=2)
    doc.add_paragraph(
        'Predict Guardian operates on a transparent consent-based model. All monitored drivers are aware '
        'of the oversight and must accept terms when their vehicle is registered. The app respects privacy '
        'while providing essential safety monitoring for families and businesses.'
    )

    privacy_features = [
        'Driver Notification: Drivers see Guardian icon indicating active monitoring',
        'Consent Required: Drivers must accept monitoring terms during profile setup',
        'Privacy Modes: Optional "privacy zones" where detailed tracking is paused',
        'Data Retention: Configurable data retention periods (7/30/90 days)',
        'Access Logs: Drivers can view who accessed their vehicle data and when'
    ]
    for feature in privacy_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('10.5 Use Case Scenarios', level=2)

    p = doc.add_paragraph()
    p.add_run('Scenario 1: Teenage Driver Safety\n').bold = True
    doc.add_paragraph(
        'A parent registers their 18-year-old son\'s vehicle under the family account. Through Guardian, '
        'they can monitor driving behavior, ensure speed limits are respected, and receive alerts if the '
        'vehicle travels outside approved areas. In case of emergency, they can instantly locate the vehicle.',
        style='List Bullet'
    )

    p = doc.add_paragraph()
    p.add_run('Scenario 2: Family Fleet Management\n').bold = True
    doc.add_paragraph(
        'A family with 5 vehicles (father, mother, two adult children, and a family driver) uses Guardian '
        'to track all vehicle maintenance schedules, ensure the elderly parent\'s vehicle is serviced on time, '
        'and monitor the hired driver\'s routes during work hours.',
        style='List Bullet'
    )

    p = doc.add_paragraph()
    p.add_run('Scenario 3: Small Business Owner\n').bold = True
    doc.add_paragraph(
        'A restaurant owner with 3 delivery vehicles uses Guardian to optimize routes, monitor driver behavior, '
        'track fuel efficiency, and ensure vehicles receive timely maintenance to avoid costly breakdowns.',
        style='List Bullet'
    )

    doc.add_page_break()

    # =========================================================================
    # Section 11: Benefits for Individual Vehicle Owners
    # =========================================================================
    doc.add_heading('11. Benefits for Individual Vehicle Owners', level=1)

    doc.add_heading('11.1 Financial Savings', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Benefit', 'Description', 'Estimated Annual Savings']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    savings_data = [
        ['Prevent Roadside Breakdowns', 'AI targets predicting failures 7-60 days ahead', 'QAR 500-2,000 per incident'],
        ['Reduce Repair Costs', 'Early detection prevents cascading damage', '30-50% reduction in repair bills'],
        ['Optimize Maintenance Timing', 'Replace parts at optimal time', 'QAR 300-800 annually'],
        ['Extend Vehicle Lifespan', 'Proactive care adds years to vehicle life', 'QAR 5,000-15,000 over lifetime']
    ]
    for i, row_data in enumerate(savings_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('11.2 Safety & Peace of Mind', level=2)
    safety_features = [
        'Battery Failure Prevention: 2-3 week advance warning before complete failure',
        'Alternator Monitoring: Detect charging system degradation early',
        'Cooling System Alerts: Prevent engine overheating in Qatar\'s extreme climate',
        'Fuel System Health: Monitor fuel pump and injector performance',
        'Transmission Monitoring: Early warning for expensive transmission issues'
    ]
    for feature in safety_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.3 Fuel Economy Monitoring', level=2)
    doc.add_paragraph(
        'Track and optimize your vehicle\'s fuel consumption with comprehensive fuel economy tools:'
    )
    fuel_features = [
        'Real-Time Fuel Economy: Monitor L/100km or MPG while driving',
        'Fuel Fillup Logging: Record fuel purchases with cost tracking',
        'Efficiency Trends: View weekly/monthly fuel economy graphs',
        'OBD vs Actual Comparison: Compare calculated vs real-world consumption',
        'Cost Analysis: Track fuel spending over time with price per km',
        'Anomaly Detection: Alert when fuel consumption deviates from normal'
    ]
    for feature in fuel_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.4 Driving Score System', level=2)
    doc.add_paragraph(
        'Get a real-time driving score (0-100) based on your driving behavior:'
    )
    driving_score_features = [
        'Smooth Acceleration Score: Rewards gentle starts (25 points)',
        'Gentle Braking Score: Rewards anticipatory braking (25 points)',
        'Speed Compliance: Monitors adherence to speed limits (25 points)',
        'Consistency Score: Rewards steady driving patterns (25 points)',
        'Trip Scores: See your score for each individual trip',
        'Historical Trends: Track improvement over time'
    ]
    for feature in driving_score_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.5 Trip Analytics', level=2)
    trip_features = [
        'Automatic Trip Detection: System detects trip start/end automatically',
        'Distance & Duration: Track every trip with precise measurements',
        'Fuel Per Trip: Calculate fuel consumption for each journey',
        'Trip History: Complete log of all trips with searchable history',
        'Route Statistics: Analyze common routes and their efficiency'
    ]
    for feature in trip_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.6 Geofencing & Safety Alerts (Qatar-Specific)', level=2)
    doc.add_paragraph(
        'GPS-based safety features designed for Qatar\'s unique conditions:'
    )
    geofence_features = [
        'Desert Area Alerts: Warning when entering remote desert zones',
        'Custom Geofences: Create alerts for specific locations (home, work, school)',
        'Entry/Exit Notifications: Know when vehicle enters or leaves zones',
        'Time-in-Zone Tracking: Monitor how long vehicle stays in each area',
        'Safety Recommendations: Climate-appropriate advice when in remote areas'
    ]
    for feature in geofence_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.7 Convenience Features', level=2)
    convenience = [
        'Mobile App Alerts: Receive predictions directly on smartphone',
        'Maintenance Scheduling: AI suggests optimal service intervals',
        'Service History: Complete digital record of all maintenance',
        'Multi-Vehicle Support: Monitor entire family fleet from one app',
        'Real-Time Dashboard: Live vehicle health status while driving',
        'VIN Decoder: Automatic vehicle identification from VIN scan',
        'DTC Code Lookup: Plain English explanations for trouble codes',
        'Data Export: Export all data to CSV/Excel for analysis'
    ]
    for feature in convenience:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.8 AI Learning Personalization', level=2)
    doc.add_paragraph('The system learns each vehicle\'s unique characteristics:')
    personalization = [
        'Baseline Learning: Establishes normal operating parameters for your specific vehicle',
        'Driving Pattern Adaptation: Adjusts predictions based on your driving style',
        'Climate Compensation: Accounts for Qatar\'s extreme heat conditions',
        'Feedback Loop: Confirms predictions to continuously improve accuracy'
    ]
    for feature in personalization:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 12: Benefits for Fleet Owners
    # =========================================================================
    doc.add_heading('12. Benefits for Fleet Owners & Enterprises', level=1)

    doc.add_heading('12.1 Operational Efficiency', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'Improvement', 'Business Impact']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    efficiency_data = [
        ['Unplanned Downtime', 'Reduce by 60-80%', 'Increased revenue, customer satisfaction'],
        ['Maintenance Costs', 'Reduce by 25-40%', 'Direct bottom-line improvement'],
        ['Vehicle Utilization', 'Increase by 10-15%', 'More revenue per vehicle'],
        ['Fuel Efficiency', 'Improve by 5-10%', 'Reduced operating costs']
    ]
    for i, row_data in enumerate(efficiency_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('12.2 Fleet Management Features', level=2)
    fleet_features = [
        'Centralized Dashboard: Monitor all vehicles from single interface',
        'Priority Alerts: Critical issues flagged for immediate attention',
        'Maintenance Planning: Schedule service based on AI predictions',
        'Driver Assignment: Consider vehicle health when assigning routes',
        'Compliance Tracking: Maintain service records for regulatory requirements',
        'Cost Analytics: Track maintenance spending per vehicle/fleet'
    ]
    for feature in fleet_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('12.3 ROI Analysis for Fleet Operations', level=2)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Example: 50-Vehicle Commercial Fleet').bold = True

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Category', 'Without System', 'With System', 'Savings']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    roi_data = [
        ['Unplanned Repairs', 'QAR 150,000', 'QAR 45,000', 'QAR 105,000'],
        ['Towing/Roadside', 'QAR 25,000', 'QAR 5,000', 'QAR 20,000'],
        ['Vehicle Downtime', 'QAR 200,000', 'QAR 60,000', 'QAR 140,000'],
        ['Preventive Maintenance', 'QAR 100,000', 'QAR 85,000', 'QAR 15,000'],
        ['Total', 'QAR 475,000', 'QAR 195,000', 'QAR 280,000']
    ]
    for i, row_data in enumerate(roi_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 4:  # Total row
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('ROI: 400%+ in first year ').bold = True
    p.add_run('(based on system cost of ~QAR 50,000 for 50-vehicle deployment)')

    doc.add_page_break()

    # =========================================================================
    # Section 13: Hardware Requirements & Pricing
    # =========================================================================
    doc.add_heading('13. Hardware Requirements & Pricing', level=1)

    doc.add_heading('13.1 Dual ESP32-S3 Sensor Unit', level=2)
    doc.add_paragraph(
        'The system uses two ESP32-S3 microcontrollers for comprehensive external sensor coverage:'
    )

    p = doc.add_paragraph()
    p.add_run('ESP32-S3 Unit #1: Engine Bay Sensors\n').bold = True
    doc.add_paragraph('Oil temperature monitoring', style='List Bullet')
    doc.add_paragraph('Vibration analysis (engine mount health)', style='List Bullet')
    doc.add_paragraph('Ambient temperature under hood', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('ESP32-S3 Unit #2: Cabin/External Sensors\n').bold = True
    doc.add_paragraph('External ambient temperature', style='List Bullet')
    doc.add_paragraph('Additional oil quality sensing', style='List Bullet')
    doc.add_paragraph('Expansion ports for custom sensors', style='List Bullet')

    doc.add_heading('13.2 Complete Hardware Bill of Materials', level=2)

    table = doc.add_table(rows=15, cols=5)
    table.style = 'Table Grid'

    headers = ['Component', 'Qty', 'Unit (USD)', 'Unit (QAR)', 'Total (QAR)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    bom_data = [
        ['ESP32-S3-WROOM-1 N8R8', '2', '$2.66', '9.70', '19.40'],
        ['ESP32-S3-DevKitC-1 Board', '2', '$8.00', '29.15', '58.30'],
        ['J1962 OBD-II 16-Pin Connector', '1', '$3.80', '13.85', '13.85'],
        ['MAX6675 + K-Type Thermocouple', '2', '$4.00', '14.58', '29.16'],
        ['MPU6050 Accelerometer/Gyro', '2', '$1.50', '5.47', '10.94'],
        ['DS18B20 Waterproof Temp Sensor', '4', '$2.00', '7.29', '29.16'],
        ['Oil Quality Sensor (Capacitive)', '1', '$15.00', '54.66', '54.66'],
        ['5V DC-DC Buck Converter', '2', '$1.50', '5.47', '10.94'],
        ['IP65 Waterproof Enclosure', '2', '$5.00', '18.22', '36.44'],
        ['Wiring Harness & Connectors', '1', '$8.00', '29.15', '29.15'],
        ['Heat-Resistant Silicone Cables', '2m', '$3.00', '10.93', '10.93'],
        ['Mounting Hardware', '1', '$5.00', '18.22', '18.22'],
        ['Shipping & Import', '-', '$20.00', '72.88', '72.88'],
        ['SUBTOTAL: ESP32 Sensor Kit', '', '', '', 'QAR 394.03']
    ]
    for i, row_data in enumerate(bom_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 13:  # Subtotal row
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i + 1].cells[j], 'E6F2FF')

    doc.add_paragraph()

    doc.add_heading('13.3 Total Hardware Cost Per Vehicle', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Configuration', 'Components', 'Total Cost (QAR)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    config_data = [
        ['Basic', 'OBD Adapter only', '43.73'],
        ['Standard', 'OBD + Single ESP32 + 3 Sensors', '220.00'],
        ['Professional', 'OBD + Dual ESP32 + Full Sensor Suite', '437.76'],
        ['Enterprise', 'Professional + OBDLink MX+', '773.00']
    ]
    for i, row_data in enumerate(config_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Note: ').italic = True
    p.add_run('Prices based on AliExpress/Alibaba bulk pricing. QAR conversion at rate: 1 USD = 3.6438 QAR')

    doc.add_page_break()

    # =========================================================================
    # Section 14: PC Hardware Recommendations
    # =========================================================================
    doc.add_heading('14. PC Hardware Recommendations', level=1)

    # Tier 1
    doc.add_heading('14.1 Tier 1: Entry Level (1-10 Vehicles)', level=2)
    p = doc.add_paragraph()
    p.add_run('Use Case: ').bold = True
    p.add_run('Home user or small business with limited vehicles')

    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'

    headers = ['Component', 'Specification', 'Price (QAR)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '228B22')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tier1_data = [
        ['Processor', 'Intel Core i5-12400 or AMD Ryzen 5 5600', '550'],
        ['RAM', '16GB DDR4 3200MHz', '180'],
        ['Storage', '512GB NVMe SSD', '150'],
        ['GPU', 'Integrated Graphics', 'Included'],
        ['Power Supply', '450W 80+ Bronze', '120'],
        ['Case', 'Mid-Tower ATX', '150'],
        ['OS', 'Windows 11 Home', '400'],
        ['Monitor', '24" 1080p IPS', '350'],
        ['TOTAL', '', 'QAR 2,100']
    ]
    for i, row_data in enumerate(tier1_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 8:
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i + 1].cells[j], 'E6FFE6')

    doc.add_paragraph()

    # Tier 2
    doc.add_heading('14.2 Tier 2: Professional (10-50 Vehicles)', level=2)
    p = doc.add_paragraph()
    p.add_run('Use Case: ').bold = True
    p.add_run('Commercial fleet operator, taxi company, delivery service')

    table = doc.add_table(rows=12, cols=3)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'FF8C00')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tier2_data = [
        ['Processor', 'Intel Core i7-13700 or AMD Ryzen 7 7700X', '1,200'],
        ['RAM', '32GB DDR5 5600MHz', '450'],
        ['Storage (Primary)', '1TB NVMe SSD (Gen4)', '350'],
        ['Storage (Data)', '2TB HDD (Backup)', '200'],
        ['GPU', 'NVIDIA RTX 3060 12GB (CUDA)', '1,100'],
        ['Power Supply', '650W 80+ Gold', '280'],
        ['Case', 'Mid-Tower with Airflow', '250'],
        ['Cooling', 'Tower CPU Cooler', '150'],
        ['OS', 'Windows 11 Pro', '600'],
        ['Monitor', '27" 1440p IPS', '700'],
        ['TOTAL', '', 'QAR 5,730']
    ]
    for i, row_data in enumerate(tier2_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 10:
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i + 1].cells[j], 'FFE6CC')

    doc.add_paragraph()

    # Tier 3
    doc.add_heading('14.3 Tier 3: Enterprise (50-500+ Vehicles)', level=2)
    p = doc.add_paragraph()
    p.add_run('Use Case: ').bold = True
    p.add_run('Major fleet operator, government vehicles, enterprise deployment')

    table = doc.add_table(rows=15, cols=3)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '8B0000')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tier3_data = [
        ['Processor', 'Intel Core i9-14900K or AMD Ryzen 9 7950X', '2,200'],
        ['RAM', '64GB DDR5 6000MHz', '900'],
        ['Storage (Primary)', '2TB NVMe SSD (Gen4)', '650'],
        ['Storage (Database)', '4TB NVMe SSD', '1,100'],
        ['Storage (Backup)', '8TB HDD RAID Array', '800'],
        ['GPU', 'NVIDIA RTX 4080 16GB', '4,200'],
        ['Power Supply', '850W 80+ Platinum', '500'],
        ['Case', 'Full Tower Server Case', '450'],
        ['Cooling', '360mm AIO Liquid Cooler', '500'],
        ['Motherboard', 'High-end Z790/X670E', '900'],
        ['OS', 'Windows 11 Pro for Workstations', '800'],
        ['Monitor', '32" 4K IPS Professional', '1,500'],
        ['UPS', '2000VA Online UPS', '1,800'],
        ['TOTAL', '', 'QAR 16,450']
    ]
    for i, row_data in enumerate(tier3_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 13:
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i + 1].cells[j], 'FFE6E6')

    doc.add_page_break()

    # =========================================================================
    # Section 15: Startup Budget Estimate
    # =========================================================================
    doc.add_heading('15. Startup Budget Estimate', level=1)

    doc.add_heading('15.1 Single Vehicle (Personal Use)', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Category', 'Item', 'Cost (QAR)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    personal_data = [
        ['Hardware', 'Predict Hardware Kit (required)', '450'],
        ['Subscription', 'Annual Subscription', '500'],
        ['Optional', 'Android Smartphone (if needed)', '500-1,500'],
        ['TOTAL YEAR 1', '', 'QAR 950 - 2,450']
    ]
    for i, row_data in enumerate(personal_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 3:
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Year 2+: QAR 500/year (renewal only)').bold = True

    doc.add_paragraph()

    doc.add_heading('15.2 Small Fleet (10 Vehicles)', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    small_fleet_data = [
        ['Hardware', 'Predict Hardware Kits (10x @ QAR 350)', '3,500'],
        ['Subscription', 'Annual Fleet Subscription (10x @ QAR 400)', '4,000'],
        ['Optional', 'Tablets for Drivers (10x)', '3,500'],
        ['Setup', 'Professional Installation + Training', '1,000'],
        ['TOTAL YEAR 1', '', 'QAR 12,000']
    ]
    for i, row_data in enumerate(small_fleet_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data
            if i == 4:
                table.rows[i + 1].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Per Vehicle Year 1: QAR 750 | Year 2+: QAR 400/vehicle (renewal only)').bold = True

    doc.add_paragraph()

    doc.add_heading('15.3 Budget Summary Table', level=2)

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    headers = ['Deployment', 'Year 1 Cost', 'Per Vehicle', 'Annual Savings', 'ROI (Year 1)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    summary_data = [
        ['Personal (1)', 'QAR 950', 'QAR 950', 'QAR 2,000-5,000', '110-425%'],
        ['Small (10)', 'QAR 7,500', 'QAR 750', 'QAR 15,000+', '100%+'],
        ['Medium (50)', 'QAR 37,500', 'QAR 750', 'QAR 75,000+', '100%+'],
        ['Enterprise (200)', 'QAR 150,000', 'QAR 750', 'QAR 300,000+', '100%+']
    ]
    for i, row_data in enumerate(summary_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_page_break()

    # =========================================================================
    # Section 16: Competitive Advantages
    # =========================================================================
    doc.add_heading('16. Competitive Advantages in Qatar & MENA', level=1)

    doc.add_heading('16.1 Market Opportunity', level=2)

    p = doc.add_paragraph()
    p.add_run('Qatar Market Size:\n').bold = True
    qatar_stats = [
        'Registered vehicles: 1.2+ million',
        'Commercial fleets: 50,000+ vehicles',
        'Annual vehicle maintenance market: QAR 2+ billion',
        'EV adoption rate: Growing 30% annually'
    ]
    for stat in qatar_stats:
        doc.add_paragraph(stat, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('MENA Region:\n').bold = True
    mena_stats = [
        'Total vehicles: 50+ million',
        'Fleet vehicles: 5+ million',
        'Combined maintenance market: $15+ billion USD'
    ]
    for stat in mena_stats:
        doc.add_paragraph(stat, style='List Bullet')

    doc.add_heading('16.2 Climate-Specific Advantages', level=2)
    doc.add_paragraph('Qatar and MENA face unique automotive challenges that Predict addresses:')

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    headers = ['Challenge', 'Predict Solution']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    climate_data = [
        ['Extreme Heat (45-50°C summers)', 'Cooling system monitoring, thermal stress prediction'],
        ['Battery Degradation', 'Accelerated in heat; AI predicts failures 3-4 weeks early'],
        ['AC System Strain', 'Monitors compressor health, refrigerant efficiency'],
        ['Dust & Sand Exposure', 'Air filter and intake system monitoring'],
        ['Stop-Start Traffic', 'Transmission and brake system wear prediction'],
        ['Long Highway Distances', 'Tire and suspension health monitoring']
    ]
    for i, row_data in enumerate(climate_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('16.3 Unique Selling Propositions', level=2)
    usps = [
        'First LSTM-Powered Consumer Solution: Deep learning AI previously only available to automotive OEMs',
        '7-60 Day Prediction Window (Target): Aiming for longest advance warning in the market',
        'Continuous Learning: System improves with every mile driven',
        'Regional Optimization: Algorithms tuned for MENA climate and driving conditions',
        'Bilingual Support: Arabic and English interface (planned)',
        'Local Support: Qatar-based technical support team',
        'Affordable Entry Point: Starting under QAR 1,500 per vehicle',
        'Open Architecture: Integrates with existing fleet management systems'
    ]
    for i, usp in enumerate(usps, 1):
        doc.add_paragraph(f'{i}. {usp}')

    doc.add_heading('16.4 Strategic Partnerships (Target)', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ['Partner Type', 'Target Organizations', 'Value Proposition']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '003366')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    partner_data = [
        ['Insurance Companies', 'Qatar Insurance, QIC, Doha Insurance', 'Reduced claims, usage-based insurance'],
        ['Fleet Operators', 'Karwa, Uber Qatar, delivery services', 'ROI from reduced downtime'],
        ['Car Dealerships', 'Al Fardan, Saleh Al Hamad, NBK', 'Extended warranty programs'],
        ['Government', 'Ministry of Transport, Qatar Post', 'Fleet optimization for public services'],
        ['Oil & Gas', 'QatarEnergy, Oryx GTL', 'Heavy vehicle fleet management']
    ]
    for i, row_data in enumerate(partner_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading('16.5 Go-to-Market Strategy for Qatar', level=2)

    p = doc.add_paragraph()
    p.add_run('Phase 1 (Months 1-6): Foundation\n').bold = True
    phase1 = [
        'Launch with 10 pilot fleet customers',
        'Establish local support center',
        'Arabic localization completion',
        'Partnership with 2-3 insurance companies'
    ]
    for item in phase1:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Phase 2 (Months 7-12): Growth\n').bold = True
    phase2 = [
        'Scale to 50 fleet customers',
        'Consumer app launch',
        'Dealership partnership program',
        'Government fleet pilot'
    ]
    for item in phase2:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Phase 3 (Year 2): Expansion\n').bold = True
    phase3 = [
        'UAE and Saudi Arabia expansion',
        '10,000+ vehicle deployments',
        'OEM partnership discussions',
        'Regional data center establishment'
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # Section 17: Future R&D Roadmap
    # =========================================================================
    doc.add_heading('17. Future R&D Roadmap', level=1)

    doc.add_paragraph(
        'Predict\'s AI-powered predictive maintenance technology is designed as a platform that can be '
        'extended beyond automotive applications. Our R&D roadmap includes strategic expansion into '
        'additional sectors where predictive maintenance can deliver significant value.'
    )

    doc.add_heading('17.1 Facility Management Predictive Maintenance', level=2)
    doc.add_paragraph(
        'Applying our AI algorithms to building systems including HVAC, elevators, electrical systems, '
        'and plumbing infrastructure. Key applications include:'
    )

    facility_apps = [
        'HVAC System Monitoring: Predict compressor failures, refrigerant leaks, and filter replacements',
        'Elevator Maintenance: Anticipate motor wear, cable fatigue, and door mechanism issues',
        'Electrical Systems: Monitor transformer health, circuit breaker performance, and power quality',
        'Plumbing Infrastructure: Detect early signs of pipe corrosion, pump degradation, and valve failures',
        'Fire Safety Systems: Predict maintenance needs for sprinklers, alarms, and suppression equipment'
    ]
    for app in facility_apps:
        doc.add_paragraph(app, style='List Bullet')

    doc.add_heading('17.2 Industrial Equipment Maintenance', level=2)
    industrial_apps = [
        'Manufacturing Machinery: CNC machines, robotic arms, conveyor systems',
        'Oil & Gas Equipment: Pumps, compressors, pipelines, drilling equipment',
        'Power Generation: Turbines, generators, transformers, solar inverters',
        'Marine Vessels: Ship engines, navigation systems, cargo handling equipment',
        'Agricultural Equipment: Tractors, harvesters, irrigation systems'
    ]
    for app in industrial_apps:
        doc.add_paragraph(app, style='List Bullet')

    doc.add_heading('17.3 Smart City Infrastructure', level=2)
    city_apps = [
        'Traffic Systems: Signal controllers, cameras, sensors',
        'Street Lighting: LED fixture health, power consumption optimization',
        'Water Treatment: Pump stations, filtration systems, chemical dosing',
        'Public Transportation: Buses, metro systems, charging infrastructure',
        'Waste Management: Collection vehicles, sorting equipment, compactors'
    ]
    for app in city_apps:
        doc.add_paragraph(app, style='List Bullet')

    doc.add_heading('17.4 R&D Timeline', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Phase', 'Focus Area', 'Timeline']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '4B0082')  # Indigo for R&D
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    rd_timeline = [
        ['Phase 1', 'Automotive Platform Maturity', 'Year 1-2'],
        ['Phase 2', 'Facility Management Pilot', 'Year 2-3'],
        ['Phase 3', 'Industrial Equipment Expansion', 'Year 3-4'],
        ['Phase 4', 'Smart City Integration', 'Year 4-5']
    ]
    for i, row_data in enumerate(rd_timeline):
        for j, cell_data in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_data

    doc.add_paragraph()
    doc.add_paragraph(
        'This expansion strategy leverages our core AI competencies in pattern recognition, anomaly detection, '
        'and remaining useful life estimation while adapting sensor integration and failure models for each domain.'
    )

    doc.add_page_break()

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('─' * 60)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    confidential = doc.add_paragraph(
        'This document contains confidential business information intended for potential '
        'investors and partners. Distribution without authorization is prohibited.'
    )
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    confidential.runs[0].italic = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Document Version: ').bold = True
    info.add_run('2.0\n')
    info.add_run('Last Updated: ').bold = True
    info.add_run(f'{datetime.now().strftime("%B %Y")}\n')
    info.add_run('Classification: ').bold = True
    info.add_run('Confidential - Investor Use Only')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_path = r'c:\D Drive\Predict\docs\EXECUTIVE_SUMMARY_INVESTORS_v9.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == '__main__':
    create_executive_summary()
